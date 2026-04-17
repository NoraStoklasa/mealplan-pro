from fastapi import APIRouter, Request, Form
from fastapi.templating import Jinja2Templates
from fastapi.responses import RedirectResponse, FileResponse, PlainTextResponse
import tempfile
from pathlib import Path
import os
from typing import List
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from recipes.recipe_database import (
    create_recipe,
    load_recipe,
    add_recipe_ingredient,
    update_recipe,
    replace_recipe_ingredients,
    delete_recipe,
)
from ingredients.database import get_all_ingredient_names
from logic.recalculate_nutrients import recalculate_nutrients
from logic.scaled_recipe import scale_recipe_to_energy
from app.labels import LABELS
from app.routes.recipes_helpers import (
    fetch_recipes_by_category,
    group_recipes_by_category,
    ingredients_from_form,
    parse_float,
    targets_by_category as parse_targets_by_category,
)
from config import RECIPE_DB_PATH

router = APIRouter(prefix="/recipes")
templates = Jinja2Templates(directory="app/templates")


# Route to list all recipes from the database.
@router.get("")
def list_recipes(request: Request):
    recipes_by_category = fetch_recipes_by_category()

    return templates.TemplateResponse(
        "recipes_list.html",
        {
            "request": request,
            "recipes_by_category": recipes_by_category,
        },
    )


# Route to show meal plan selection (export).
@router.get("/meal-plan")
def meal_plan(request: Request):
    recipes_by_category = fetch_recipes_by_category()

    return templates.TemplateResponse(
        "meal_plan.html",
        {
            "request": request,
            "recipes_by_category": recipes_by_category,
        },
    )


# Route to show the form for adding a new recipe.
@router.get("/new")
def new_recipe_form(request: Request):
    return templates.TemplateResponse(
        "recipe_form.html",
        {"request": request},
    )


# Route to handle the submission of a new recipe.
@router.post("/new")
def create_recipe_post(
    name: str = Form(...),
    category: str = Form(...),
    instructions: str = Form(...),
    image_path: str = Form(""),
):
    recipe_id = create_recipe(
        name=name,
        category=category,
        instructions=instructions,
        image_path=image_path,
    )

    return RedirectResponse(
        url=f"/recipes/{recipe_id}/edit",
        status_code=303,
    )


@router.post("/export")
def export_recipes(request: Request, recipe_ids: str = Form("")):
    ids = []
    for rid in recipe_ids.split(","):
        rid = rid.strip()
        if not rid:
            continue
        try:
            ids.append(int(rid))
        except ValueError:
            continue
    recipes = []
    for rid in ids:
        recipe = load_recipe(rid)
        if not recipe:
            continue
        nutrients = recalculate_nutrients(recipe)
        recipes.append(
            {
                "id": rid,
                "energy_kj": nutrients.get("energy_kj", 0) or 0,
                **recipe,
            }
        )
    recipes_by_category = group_recipes_by_category(
        recipes, lambda r: r["category"], lambda r: r["name"]
    )

    return templates.TemplateResponse(
        "recipes_export.html",
        {
            "request": request,
            "recipes_by_category": recipes_by_category,
            "lang": "en",
        },
    )


@router.post("/export/scale")
async def scale_selected_recipes(request: Request):
    form = await request.form()
    recipe_ids = form.getlist("recipe_ids")
    client_name = form.get("client_name") or ""
    lang = form.get("lang") or "en"
    labels = LABELS.get(lang, LABELS["en"])
    targets_by_category = parse_targets_by_category(form)
    recipes = []
    for rid_value in recipe_ids:
        try:
            rid = int(rid_value)
        except ValueError:
            continue
        recipe = load_recipe(rid)
        if recipe:
            recipes.append((rid, recipe))

    scaled_recipes = []
    for rid, recipe in recipes:
        category = recipe.get("category", "")
        recipe_target = parse_float(form.get(f"target_kj_recipe_{rid}"))
        target_kj = recipe_target or (targets_by_category.get(category, 0) or 0)
        if target_kj > 0:
            scaled_recipe, scaled_nutrients = scale_recipe_to_energy(recipe, target_kj)
        else:
            scaled_recipe = recipe
            scaled_nutrients = recalculate_nutrients(recipe)
        scaled_recipes.append(
            {
                "recipe_id": rid,
                "recipe": scaled_recipe,
                "nutrients": scaled_nutrients,
                "category": category,
            }
        )
    return templates.TemplateResponse(
        "recipes_scaled_preview.html",
        {
            "request": request,
            "scaled_recipes": scaled_recipes,
            "targets_by_category": targets_by_category,
            "client_name": client_name,
            "labels": labels,
            "lang": lang,
            "error": None,
        },
    )


@router.post("/export/preview-update")
async def update_preview(request: Request):
    form = await request.form()
    recipe_ids = form.getlist("recipe_ids")
    client_name = form.get("client_name") or ""
    lang = form.get("lang") or "en"
    labels = LABELS.get(lang, LABELS["en"])
    targets_by_category = parse_targets_by_category(form)

    scaled_recipes = []
    for rid_value in recipe_ids:
        try:
            rid = int(rid_value)
        except ValueError:
            continue
        recipe = load_recipe(rid)
        if not recipe:
            continue
        recipe["ingredients"] = ingredients_from_form(form, rid)
        nutrients = recalculate_nutrients(recipe)
        scaled_recipes.append(
            {
                "recipe_id": rid,
                "recipe": recipe,
                "nutrients": nutrients,
                "category": recipe.get("category", ""),
            }
        )

    return templates.TemplateResponse(
        "recipes_scaled_preview.html",
        {
            "request": request,
            "scaled_recipes": scaled_recipes,
            "targets_by_category": targets_by_category,
            "client_name": client_name,
            "labels": labels,
            "lang": lang,
            "error": None,
        },
    )


@router.post("/export/docx")
async def export_docx(request: Request):
    try:
        from docx import Document
    except ModuleNotFoundError:
        return PlainTextResponse(
            "python-docx is not installed. Run: pip install python-docx",
            status_code=500,
        )

    form = await request.form()
    recipe_ids = form.getlist("recipe_ids")
    client_name = (form.get("client_name") or "").strip()
    lang = form.get("lang") or "en"
    labels = LABELS.get(lang, LABELS["en"])

    recipes = []
    for rid_value in recipe_ids:
        try:
            rid = int(rid_value)
        except ValueError:
            continue
        recipe = load_recipe(rid)
        if not recipe:
            continue
        recipe["ingredients"] = ingredients_from_form(form, rid)
        nutrients = recalculate_nutrients(recipe)
        recipes.append({"recipe": recipe, "nutrients": nutrients})

    doc = Document()

    # Add very light beige background to entire document
    background = OxmlElement('w:background')
    background.set(qn('w:color'), '#FEFCF7')
    doc.element.insert(0, background)

    # Add logo to header (appears on all pages)
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    logo_path = Path("app/logos/logo.png")
    if logo_path.exists():
        try:
            logo_run = header_para.add_run()
            logo_run.add_picture(str(logo_path), width=Inches(0.8))
        except Exception:
            pass

    if client_name:
        client_heading = doc.add_heading(client_name, level=1)
        # Set client name heading to dark purple
        for run in client_heading.runs:
            run.font.color.rgb = RGBColor(75, 0, 130)
        client_heading.paragraph_format.space_before = 0
        client_heading.paragraph_format.space_after = 18
        client_heading.paragraph_format.line_spacing = 1.0

    # Group recipes by meal type
    meal_types = {"breakfast": [], "lunch": [], "dinner": [], "snack": []}
    for item in recipes:
        category = item["recipe"].get("category", "").lower()
        if category in meal_types:
            meal_types[category].append(item)

    # Display recipes grouped by meal type with headers
    meal_type_labels = {
        "breakfast": labels.get("breakfast", "Breakfast"),
        "lunch": labels.get("lunch", "Lunch"),
        "dinner": labels.get("dinner", "Dinner"),
        "snack": labels.get("snack", "Snacks"),
    }

    for meal_type in ["breakfast", "lunch", "dinner", "snack"]:
        if meal_types[meal_type]:
            heading = doc.add_heading(meal_type_labels[meal_type], level=0)
            # Set heading to dark purple and increase size
            for run in heading.runs:
                run.font.color.rgb = RGBColor(75, 0, 130)
                run.font.size = Pt(24)
            heading.paragraph_format.space_before = 12
            heading.paragraph_format.space_after = 12
            heading.paragraph_format.line_spacing = 1.0
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add purple bottom border to heading
            pPr = heading._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '24')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '4B0082')
            pBdr.append(bottom)
            pPr.append(pBdr)

            for item in meal_types[meal_type]:
                recipe = item["recipe"]
                nutrients = item["nutrients"]

                # Create a 2-column table with 1 row
                table = doc.add_table(rows=1, cols=2)
                table.autofit = False
                table.allow_autofit = False

                # Left cell for recipe content
                left_cell = table.rows[0].cells[0]
                right_cell = table.rows[0].cells[1]

                # Add recipe name in left cell
                recipe_name_para = left_cell.paragraphs[0]
                recipe_name_run = recipe_name_para.add_run(recipe.get("name", "Recipe"))
                recipe_name_run.bold = True
                recipe_name_run.font.color.rgb = RGBColor(75, 0, 130)
                recipe_name_run.font.size = Pt(13)
                recipe_name_para.paragraph_format.space_before = Pt(5)
                recipe_name_para.paragraph_format.space_after = Pt(5)
                recipe_name_para.paragraph_format.line_spacing = 1.0

                # Remove spacing before table by accessing the paragraph before it
                tbl_element = table._element
                p_before = tbl_element.getprevious()
                if p_before is not None:
                    pPr = p_before.get_or_add_pPr()
                    spacing = pPr.find(qn('w:spacing'))
                    if spacing is not None:
                        spacing.set(qn('w:after'), '0')
                    else:
                        spacing = OxmlElement('w:spacing')
                        spacing.set(qn('w:after'), '0')
                        spacing.set(qn('w:line'), '240')
                        spacing.set(qn('w:lineRule'), 'auto')
                        pPr.append(spacing)

                # Remove table borders (make lines invisible)
                tbl = table._element
                tblPr = tbl.tblPr
                if tblPr is None:
                    tblPr = OxmlElement('w:tblPr')
                    tbl.insert(0, tblPr)

                # Add table borders element with no borders
                tblBorders = OxmlElement('w:tblBorders')
                for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'none')
                    border.set(qn('w:sz'), '0')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), 'auto')
                    tblBorders.append(border)

                # Remove existing borders if any
                for existing in tblPr.findall(qn('w:tblBorders')):
                    tblPr.remove(existing)
                tblPr.append(tblBorders)

                # Set minimal cell margins
                tblCellMar = OxmlElement('w:tblCellMar')
                for margin_name in ['top', 'left', 'bottom', 'right']:
                    margin = OxmlElement(f'w:{margin_name}')
                    margin.set(qn('w:w'), '10')
                    margin.set(qn('w:type'), 'dxa')
                    tblCellMar.append(margin)

                for existing in tblPr.findall(qn('w:tblCellMar')):
                    tblPr.remove(existing)
                tblPr.append(tblCellMar)

                # Set row height
                for row in table.rows:
                    row.height = Inches(0.3)

                # Set column widths
                table.columns[0].width = Inches(3.5)
                table.columns[1].width = Inches(2)

                # Left cell for recipe content is already set above as table.rows[1].cells[0]
                # Add ingredients
                left_para = left_cell.add_paragraph(f"{labels['ingredients']}:")
                left_para.style = 'Normal'
                left_para.paragraph_format.space_before = 0
                left_para.paragraph_format.space_after = 0
                left_para.paragraph_format.line_spacing = 1.0
                # Make "Ingredients:" bold
                ing_run = left_para.runs[0]
                ing_run.bold = True
                for ing in recipe.get("ingredients", []):
                    ing_para = left_cell.add_paragraph(
                        f"{ing.get('name', '')} – {ing.get('portion_g', 0)} g",
                        style="List Bullet",
                    )
                    ing_para.paragraph_format.space_before = Pt(5)
                    ing_para.paragraph_format.space_after = Pt(3)
                    ing_para.paragraph_format.line_spacing = 1.0

                # Add instructions
                instructions = recipe.get("instructions") or ""
                if instructions:
                    instr_para = left_cell.add_paragraph(f"{labels['instructions']}:")
                    instr_para.paragraph_format.space_before = Pt(5)
                    instr_para.paragraph_format.space_after = Pt(3)
                    instr_para.paragraph_format.line_spacing = 1.0
                    instr_para.runs[0].bold = True
                    instr_content = left_cell.add_paragraph(instructions)
                    instr_content.paragraph_format.space_before = 0
                    instr_content.paragraph_format.space_after = 0
                    instr_content.paragraph_format.line_spacing = 1.0

                # Add nutrition summary
                nutrition_para = left_cell.add_paragraph(f"{labels['nutrition_summary']}:")
                nutrition_para.paragraph_format.space_before = Pt(5)
                nutrition_para.paragraph_format.space_after = Pt(3)
                nutrition_para.paragraph_format.line_spacing = 1.0
                nutrition_para.runs[0].bold = True

                energy_para = left_cell.add_paragraph(f"{labels['energy']}: {nutrients['energy_kj']} kJ")
                energy_para.paragraph_format.space_before = 0
                energy_para.paragraph_format.space_after = 0
                energy_para.paragraph_format.line_spacing = 1.0

                protein_para = left_cell.add_paragraph(f"{labels['protein']}: {nutrients['protein_g']} g")
                protein_para.paragraph_format.space_before = 0
                protein_para.paragraph_format.space_after = 0
                protein_para.paragraph_format.line_spacing = 1.0

                carbs_para = left_cell.add_paragraph(f"{labels['carbs']}: {nutrients['carbs_g']} g")
                carbs_para.paragraph_format.space_before = 0
                carbs_para.paragraph_format.space_after = 0
                carbs_para.paragraph_format.line_spacing = 1.0

                fat_para = left_cell.add_paragraph(f"{labels['fat']}: {nutrients['fat_g']} g")
                fat_para.paragraph_format.space_before = 0
                fat_para.paragraph_format.space_after = 0
                fat_para.paragraph_format.line_spacing = 1.0

                fibre_para = left_cell.add_paragraph(f"{labels['fibre']}: {nutrients['fibre_g']} g")
                fibre_para.paragraph_format.space_before = 0
                fibre_para.paragraph_format.space_after = 0
                fibre_para.paragraph_format.line_spacing = 1.0

                # Right cell for image
                right_cell = table.rows[0].cells[1]

                # Try to find recipe-specific image, fall back to cheesecake.jpg
                project_dir = Path(__file__).parent.parent.parent
                recipe_name = recipe.get("name", "").replace(" ", "_").lower()
                recipe_image_path = project_dir / f"images/{recipe_name}.jpg"
                default_image_path = project_dir / "images/cheesecake.jpg"

                image_to_use = None
                if recipe_image_path.exists():
                    image_to_use = recipe_image_path
                elif default_image_path.exists():
                    image_to_use = default_image_path

                if image_to_use:
                    try:
                        image_para = right_cell.add_paragraph()
                        image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = image_para.add_run()
                        run.add_picture(str(image_to_use), width=Inches(1.8))
                    except Exception:
                        pass

                # Add spacing after each recipe
                spacing_para = doc.add_paragraph()
                spacing_para.paragraph_format.space_before = 12
                spacing_para.paragraph_format.space_after = 0

    # Add page numbers to footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Create page number field
    run = footer_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        temp_path = tmp.name
    doc.save(temp_path)

    # Create filename from client name or use default
    if client_name:
        filename = client_name.replace(" ", "_") + "_mealplan.docx"
    else:
        filename = "mealplan_pro_export.docx"

    return FileResponse(
        temp_path,
        filename=filename,
        media_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )


# Route to show the details of a specific recipe.
@router.get("/{recipe_id}")
def recipe_detail(request: Request, recipe_id: int):
    recipe = load_recipe(recipe_id)
    if not recipe:
        return RedirectResponse(url="/recipes", status_code=303)
    ingredients = get_all_ingredient_names()

    nutrients = recalculate_nutrients(recipe)

    return templates.TemplateResponse(
        "recipe_detail.html",
        {
            "request": request,
            "recipe_id": recipe_id,
            "recipe": recipe,
            "ingredients": ingredients,
            "nutrients": nutrients,
        },
    )


@router.get("/{recipe_id}/edit")
def edit_recipe_form(request: Request, recipe_id: int):
    recipe = load_recipe(recipe_id)
    if not recipe:
        return RedirectResponse(url="/recipes", status_code=303)
    return templates.TemplateResponse(
        "recipe_edit.html",
        {
            "request": request,
            "recipe_id": recipe_id,
            "recipe": recipe,
        },
    )


# Route to handle adding an ingredient to a specific recipe
@router.post("/{recipe_id}")
def add_ingredient_to_recipe(
    recipe_id: int,
    ingredient_name: str = Form(...),
    portion_g: float = Form(...),
):
    add_recipe_ingredient(
        recipe_id=recipe_id,
        ingredient_name=ingredient_name,
        portion_g=portion_g,
    )

    return RedirectResponse(
        url=f"/recipes/{recipe_id}",
        status_code=303,
    )


@router.post("/{recipe_id}/edit")
def update_recipe_post(
    recipe_id: int,
    name: str = Form(...),
    category: str = Form(...),
    instructions: str = Form(...),
    image_path: str = Form(""),
    ingredient_name: List[str] = Form([]),
    portion_g: List[str] = Form([]),
):
    update_recipe(
        recipe_id=recipe_id,
        name=name,
        category=category,
        instructions=instructions,
        image_path=image_path,
    )
    ingredients = []
    for name_value, portion_value in zip(ingredient_name, portion_g):
        name_value = (name_value or "").strip()
        if not name_value:
            continue
        try:
            portion = float(portion_value)
        except (TypeError, ValueError):
            continue
        if portion <= 0:
            continue
        ingredients.append({"name": name_value, "portion_g": portion})
    replace_recipe_ingredients(recipe_id, ingredients)

    return RedirectResponse(
        url=f"/recipes/{recipe_id}",
        status_code=303,
    )


@router.post("/{recipe_id}/delete")
def delete_recipe_post(recipe_id: int):
    delete_recipe(recipe_id)
    return RedirectResponse(url="/recipes", status_code=303)
